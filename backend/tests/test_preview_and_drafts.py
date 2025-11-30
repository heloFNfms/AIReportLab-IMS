import os
import io
import unittest
from datetime import datetime
from fastapi.testclient import TestClient
from app.main import app
from app.core.config import settings
from app.core.deps import get_current_user
from app.services.ai.template_analyzer import TemplateAnalyzer


class DummyUser:
    def __init__(self, id: int):
        self.id = id
        self.username = "tester"


def override_current_user():
    return DummyUser(1)


class PreviewAndDraftsTestCase(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        app.dependency_overrides[get_current_user] = override_current_user
        os.makedirs(settings.UPLOAD_FOLDER, exist_ok=True)

    @classmethod
    def tearDownClass(cls):
        app.dependency_overrides.clear()

    def test_config_endpoint(self):
        client = TestClient(app)
        resp = client.get("/api/config")
        self.assertEqual(resp.status_code, 200)
        data = resp.json()
        self.assertIn("featureFlags", data)
        self.assertIn("disableAIGeneration", data["featureFlags"])

    def test_drafts_save_and_get(self):
        client = TestClient(app)
        payload = {
            "template_id": 123,
            "content": "# 草稿测试\n内容...",
            "format": "markdown"
        }
        r1 = client.post("/api/reports/drafts", json=payload)
        self.assertEqual(r1.status_code, 200)
        r2 = client.get("/api/reports/drafts/123")
        self.assertEqual(r2.status_code, 200)
        self.assertIn("content", r2.json())

    def test_template_analyzer_extract_text(self):
        # txt
        base = os.path.join(settings.UPLOAD_FOLDER, "tests")
        os.makedirs(base, exist_ok=True)
        txt_path = os.path.join(base, "sample.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("hello world")
        ta = TemplateAnalyzer()
        self.assertEqual(ta.extract_text_from_txt(txt_path), "hello world")

        # docx
        import docx
        docx_path = os.path.join(base, "sample.docx")
        d = docx.Document()
        d.add_paragraph("第一段")
        d.add_paragraph("第二段")
        d.save(docx_path)
        text = ta.extract_text_from_docx(docx_path)
        self.assertTrue("第一段" in text and "第二段" in text)


if __name__ == "__main__":
    unittest.main()
